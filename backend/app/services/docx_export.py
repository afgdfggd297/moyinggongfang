"""DOCX 导出服务 - Markdown → DOCX"""
import re
import logging
from pathlib import Path
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.core.config import get_settings

settings = get_settings()
logger = logging.getLogger(__name__)


class DocxExportService:
    """Markdown → DOCX 导出服务"""

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        # 设置正文样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.5

    def _add_heading(self, doc: Document, text: str, level: int):
        """添加标题"""
        heading = doc.add_heading(text, level=level)
        # 设置中文字体
        for run in heading.runs:
            run.font.name = '黑体'

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False, italic: bool = False):
        """添加段落"""
        p = doc.add_paragraph()
        # 处理粗体和斜体
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(12)
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.name = '宋体'
                run.font.size = Pt(12)
            else:
                run = p.add_run(part)
                run.bold = bold
                run.italic = italic
                run.font.name = '宋体'
                run.font.size = Pt(12)

    def _add_list(self, doc: Document, items: list, ordered: bool = False):
        """添加列表"""
        style_name = 'List Number' if ordered else 'List Bullet'
        for i, item in enumerate(items):
            if ordered:
                p = doc.add_paragraph(f"{i+1}. {item}")
            else:
                p = doc.add_paragraph(f"• {item}")
            p.style = doc.styles[style_name]

    def _add_table(self, doc: Document, headers: list, rows: list):
        """添加表格"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # 表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # 设置表头样式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        # 数据行
        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                row.cells[i].text = str(cell_text)

    def _add_blockquote(self, doc: Document, text: str):
        """添加引用"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.border_left = Pt(2)
        run = p.add_run(text)
        run.italic = True
        run.font.name = '宋体'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _parse_markdown_content(self, md_content: str) -> list:
        """解析 Markdown 内容为结构化数据"""
        lines = md_content.split('\n')
        elements = []
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # 空行
            if not line:
                i += 1
                continue

            # 标题
            if line.startswith('#'):
                match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if match:
                    level = len(match.group(1))
                    text = match.group(2).strip()
                    elements.append({
                        'type': 'heading',
                        'level': level,
                        'text': text
                    })
                i += 1
                continue

            # 表格
            if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                headers = [h.strip() for h in line.split('|') if h.strip()]
                rows = []
                i += 2  # 跳过表头和分隔行
                while i < len(lines) and '|' in lines[i]:
                    row = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
                    if row:
                        rows.append(row)
                    i += 1
                elements.append({
                    'type': 'table',
                    'headers': headers,
                    'rows': rows
                })
                continue

            # 引用
            if line.startswith('>'):
                quote_text = line[1:].strip()
                elements.append({
                    'type': 'blockquote',
                    'text': quote_text
                })
                i += 1
                continue

            # 无序列表
            if line.startswith('- ') or line.startswith('* '):
                items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                    items.append(lines[i].strip()[2:])
                    i += 1
                elements.append({
                    'type': 'list',
                    'ordered': False,
                    'items': items
                })
                continue

            # 有序列表
            if re.match(r'^\d+\.\s', line):
                items = []
                while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                    items.append(re.sub(r'^\d+\.\s', '', lines[i].strip()))
                    i += 1
                elements.append({
                    'type': 'list',
                    'ordered': True,
                    'items': items
                })
                continue

            # 分隔线
            if line.startswith('---') or line.startswith('***') or line.startswith('___'):
                elements.append({
                    'type': 'hr'
                })
                i += 1
                continue

            # 普通段落
            para_text = line
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#') and not lines[i].strip().startswith('- ') and not lines[i].strip().startswith('* ') and not lines[i].strip().startswith('>') and not re.match(r'^\d+\.\s', lines[i].strip()) and '|' not in lines[i]:
                para_text += ' ' + lines[i].strip()
                i += 1
            elements.append({
                'type': 'paragraph',
                'text': para_text
            })

        return elements

    def markdown_to_docx(self, md_content: str, title: str, plan_id: str) -> str:
        """将 Markdown 内容转换为 DOCX 文件"""
        try:
            doc = Document()
            self._setup_styles(doc)

            # 添加文档标题
            doc.add_heading(title, level=0)

            # 解析 Markdown 内容
            elements = self._parse_markdown_content(md_content)

            for elem in elements:
                elem_type = elem.get('type')

                if elem_type == 'heading':
                    level = elem.get('level', 1)
                    # DOCX 支持 0-9 级标题，但通常只用 1-4 级
                    if level > 4:
                        level = 4
                    self._add_heading(doc, elem['text'], level)
                elif elem_type == 'paragraph':
                    text = elem.get('text', '')
                    if text:
                        self._add_paragraph(doc, text)
                elif elem_type == 'list':
                    items = elem.get('items', [])
                    if items:
                        self._add_list(doc, items, elem.get('ordered', False))
                elif elem_type == 'table':
                    headers = elem.get('headers', [])
                    rows = elem.get('rows', [])
                    if headers and rows:
                        self._add_table(doc, headers, rows)
                elif elem_type == 'blockquote':
                    text = elem.get('text', '')
                    if text:
                        self._add_blockquote(doc, text)
                elif elem_type == 'hr':
                    doc.add_paragraph('─' * 50)

            # 保存文件
            filename = f"docx_{plan_id}.docx"
            filepath = self.upload_dir / filename
            doc.save(str(filepath))

            logger.info("[docx_export] 导出成功: %s", filepath)
            return str(filepath)

        except Exception as e:
            logger.error("[docx_export] 导出失败: %s", str(e))
            raise

    def markdown_to_docx_sync(self, md_content: str, title: str, plan_id: str) -> str:
        """同步版本的导出方法"""
        return self.markdown_to_docx(md_content, title, plan_id)


# 全局实例
docx_export_service = DocxExportService()
